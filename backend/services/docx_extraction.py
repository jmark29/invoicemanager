"""DOCX extraction service for importing existing client invoices.

Parses .docx files matching the 29ventures invoice template format.
Extracts: invoice number, date, service period, client, line items, totals.
"""

import re
from dataclasses import dataclass, field
from datetime import date, datetime
from pathlib import Path

from docx import Document


@dataclass
class ExtractedLineItem:
    position: int
    description: str
    amount: float


@dataclass
class ExtractedClientInvoiceData:
    invoice_number: str | None = None
    invoice_date: date | None = None
    period_start: date | None = None
    period_end: date | None = None
    client_name: str | None = None
    line_items: list[ExtractedLineItem] = field(default_factory=list)
    net_total: float | None = None
    tax_rate: float | None = None
    tax_amount: float | None = None
    gross_total: float | None = None
    raw_text: str = ""
    confidence: str = "high"


def _parse_german_amount(text: str) -> float | None:
    """Parse German-formatted amount like '16.450,00 €' → 16450.00."""
    cleaned = text.replace("€", "").replace("EUR", "").strip()
    if not cleaned:
        return None
    # Remove thousand separator (period), replace decimal comma with period
    cleaned = cleaned.replace(".", "").replace(",", ".")
    try:
        return float(cleaned)
    except ValueError:
        return None


def _parse_german_date(text: str) -> date | None:
    """Parse German date like '28.02.2025' → date(2025, 2, 28)."""
    try:
        return datetime.strptime(text.strip(), "%d.%m.%Y").date()
    except ValueError:
        return None


def extract_client_invoice_docx(file_path: str | Path) -> ExtractedClientInvoiceData:
    """Extract structured invoice data from a .docx file.

    Designed for the 29ventures invoice template with:
    - Recipient block in first paragraphs
    - "Wiesbaden, DD.MM.YYYY" date line
    - "Rechnung YYYYMM-NN" invoice number
    - "Leistungszeitraum DD.MM.YYYY bis DD.MM.YYYY" period
    - Table with Pos/Bezeichnung/Betrag columns + totals
    """
    doc = Document(str(file_path))
    result = ExtractedClientInvoiceData()

    # Collect all non-empty paragraph texts
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    result.raw_text = "\n".join(paragraphs)

    # -- Extract client name (first non-empty paragraph, excluding sender line) --
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        # First paragraph usually has sender line + client name
        # Client name is the main text (without the small sender line)
        lines = text.split("\n")
        for line in lines:
            line = line.strip()
            if line and "GmbH" not in line and "Kleist" not in line and "Wiesbaden" not in line:
                if not result.client_name:
                    result.client_name = line
                    break
        if result.client_name:
            break

    # -- Extract date, invoice number, period from paragraphs --
    for text in paragraphs:
        # Date: "Wiesbaden, 28.02.2025"
        date_match = re.search(r"Wiesbaden,\s+(\d{2}\.\d{2}\.\d{4})", text)
        if date_match and not result.invoice_date:
            result.invoice_date = _parse_german_date(date_match.group(1))

        # Invoice number: "Rechnung 202501-02"
        inv_match = re.search(r"Rechnung\s+(\d{6}-\d{2})", text)
        if inv_match and not result.invoice_number:
            result.invoice_number = inv_match.group(1)

        # Period: "Leistungszeitraum 01.01.2025 bis 31.01.2025"
        period_match = re.search(
            r"Leistungszeitraum\s+(\d{2}\.\d{2}\.\d{4})\s+bis\s+(\d{2}\.\d{2}\.\d{4})",
            text,
        )
        if period_match and not result.period_start:
            result.period_start = _parse_german_date(period_match.group(1))
            result.period_end = _parse_german_date(period_match.group(2))

    # -- Extract line items and totals from tables --
    for table in doc.tables:
        if len(table.rows) < 2:
            continue

        # Check if this is the invoice table (first row should have "Pos" and "Betrag")
        header_cells = [cell.text.strip().lower() for cell in table.rows[0].cells]
        if "pos" not in header_cells or "betrag" not in header_cells:
            continue

        for row in table.rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]
            pos_text = cells[0] if cells else ""
            # Bezeichnung is in column 1 (or 2 — merged cells duplicate)
            description = cells[1] if len(cells) > 1 else ""
            # Betrag is in the last column
            amount_text = cells[-1] if cells else ""

            if pos_text and pos_text.isdigit():
                # This is a line item row
                amount = _parse_german_amount(amount_text)
                if amount is not None:
                    result.line_items.append(
                        ExtractedLineItem(
                            position=int(pos_text),
                            description=description,
                            amount=amount,
                        )
                    )
            elif not pos_text:
                # Totals rows (Netto, USt, Brutto) — label is in cells[2]
                label = cells[2] if len(cells) > 2 else description
                if not label:
                    continue
                amount = _parse_german_amount(amount_text)
                if amount is None:
                    continue

                if "netto" in label.lower():
                    result.net_total = amount
                elif "umsatzsteuer" in label.lower():
                    result.tax_amount = amount
                    # Extract tax rate from "Umsatzsteuer 19%"
                    rate_match = re.search(r"(\d+)%", label)
                    if rate_match:
                        result.tax_rate = float(rate_match.group(1))
                elif "brutto" in label.lower():
                    result.gross_total = amount

    # Set confidence based on how much data was extracted
    if all([
        result.invoice_number,
        result.invoice_date,
        result.line_items,
        result.net_total,
    ]):
        result.confidence = "high"
    elif result.line_items or result.net_total:
        result.confidence = "medium"
    else:
        result.confidence = "low"

    return result


def extract_client_invoice_pdf(file_path: str | Path) -> ExtractedClientInvoiceData:
    """Best-effort extraction from a PDF invoice using pdfplumber.

    Falls back gracefully — returns whatever can be parsed, with low confidence.
    """
    import pdfplumber

    result = ExtractedClientInvoiceData(confidence="low")

    try:
        with pdfplumber.open(str(file_path)) as pdf:
            full_text = ""
            for page in pdf.pages:
                text = page.extract_text() or ""
                full_text += text + "\n"

            result.raw_text = full_text

            # Try the same regex patterns as DOCX
            inv_match = re.search(r"Rechnung\s+(\d{6}-\d{2})", full_text)
            if inv_match:
                result.invoice_number = inv_match.group(1)

            date_match = re.search(r"Wiesbaden,\s+(\d{2}\.\d{2}\.\d{4})", full_text)
            if date_match:
                result.invoice_date = _parse_german_date(date_match.group(1))

            period_match = re.search(
                r"Leistungszeitraum\s+(\d{2}\.\d{2}\.\d{4})\s+bis\s+(\d{2}\.\d{2}\.\d{4})",
                full_text,
            )
            if period_match:
                result.period_start = _parse_german_date(period_match.group(1))
                result.period_end = _parse_german_date(period_match.group(2))

            # Try to extract line items from tables
            for page in pdf.pages:
                tables = page.extract_tables()
                for table in tables:
                    for row in table:
                        if not row or not row[0]:
                            continue
                        pos_text = str(row[0]).strip()
                        if pos_text.isdigit() and len(row) >= 3:
                            desc = str(row[1]).strip() if row[1] else ""
                            amt_text = str(row[-1]).strip() if row[-1] else ""
                            amount = _parse_german_amount(amt_text)
                            if amount is not None and desc:
                                result.line_items.append(
                                    ExtractedLineItem(
                                        position=int(pos_text),
                                        description=desc,
                                        amount=amount,
                                    )
                                )

            # Update confidence
            if result.invoice_number and result.line_items:
                result.confidence = "medium"

    except Exception:
        pass  # Best-effort — return what we have

    return result
